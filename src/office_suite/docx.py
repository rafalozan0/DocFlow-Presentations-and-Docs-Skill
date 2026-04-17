import os
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_word(title: str, content: str, output_path: str, **kwargs) -> Dict[str, Any]:
    """Create a DOCX file from simple markdown-like text."""
    doc = Document()

    doc.styles["Normal"].font.name = kwargs.get("font_name", "Calibri")
    doc.styles["Normal"].font.size = Pt(kwargs.get("font_size", 11))

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    lines = content.split("\n")
    for line in lines:
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            para = doc.add_heading(line[2:], level=1)
            para.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        elif line.startswith("## "):
            para = doc.add_heading(line[3:], level=2)
            para.runs[0].font.color.rgb = RGBColor(51, 153, 255)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(line[2:])
        elif line[:3].isdigit() and line[1:3] == ". ":
            para = doc.add_paragraph(style="List Number")
            para.add_run(line)
        else:
            doc.add_paragraph(line)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    return {
        "output_path": output_path,
        "file_size": os.path.getsize(output_path),
        "paragraphs": len(doc.paragraphs),
    }


def add_watermark(
    input_path: str,
    watermark_text: str,
    output_path: Optional[str] = None,
    **kwargs,
) -> Dict[str, Any]:
    """Add a simple text watermark in the document header."""
    output_path = output_path or input_path

    doc = Document(input_path)
    for section in doc.sections:
        header = section.header
        para = header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(watermark_text)
        run.font.size = Pt(kwargs.get("size", 36))
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.bold = True

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return {"output_path": output_path}


def extract_text(input_path: str, **kwargs) -> str:
    """Extract plain text from a DOCX file."""
    doc = Document(input_path)
    return "\n".join(para.text for para in doc.paragraphs)


def merge_documents(input_paths, output_path: str, **kwargs) -> Dict[str, Any]:
    """Merge multiple DOCX files into one document."""
    merged_doc = Document()

    appended = 0
    for i, path in enumerate(input_paths):
        if not os.path.exists(path):
            continue

        if i > 0:
            merged_doc.add_section(WD_SECTION.NEW_PAGE)

        source_doc = Document(path)
        for element in source_doc.element.body:
            merged_doc.element.body.append(element)
        appended += 1

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    merged_doc.save(output_path)
    return {"output_path": output_path, "merged_count": appended}
